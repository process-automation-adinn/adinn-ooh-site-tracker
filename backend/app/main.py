from datetime import date, datetime
import os
import json
from pathlib import Path
import shutil
import re
from urllib.parse import quote
from typing import Optional
from uuid import uuid4

from fastapi import Depends, FastAPI, File, Form, HTTPException, UploadFile, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, Response, StreamingResponse
from fastapi.staticfiles import StaticFiles
from openpyxl import Workbook
from docx import Document
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
import io
from sqlalchemy import inspect, text as sql_text
from sqlalchemy.exc import IntegrityError
from sqlalchemy.orm import Session

from .auth import create_token, get_current_user, hash_password, require_admin, verify_password
from .database import engine, get_db, init_database, ping_database
from .models import OOHSite, StoredFile, User
from .schemas import CrudPermissionUpdate, LoginRequest, SiteOut, TokenResponse, UserCreate, UserOut, UserUpdate

BASE_DIR = Path(__file__).resolve().parent.parent
UPLOAD_DIR = Path(os.getenv("UPLOAD_DIR") or (BASE_DIR / "app" / "uploads"))
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

app = FastAPI(title="ADINN OOH Site Tracker API", version="3.1.0-neon-postgres-agreement")


def get_allowed_origins() -> list[str]:
    configured = os.getenv("CORS_ORIGINS", "").strip()
    if configured:
        return [origin.strip() for origin in configured.split(",") if origin.strip()]
    return [
        "http://localhost:5173",
        "http://127.0.0.1:5173",
        "http://localhost:3000",
        "http://127.0.0.1:3000",
    ]


app.add_middleware(
    CORSMiddleware,
    allow_origins=get_allowed_origins(),
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Static uploads are kept only as a fallback for older local files.
app.mount("/uploads", StaticFiles(directory=UPLOAD_DIR), name="uploads")

FILE_FIELD_LABELS = {
    "site_photo_file": "Site Photo",
    "landlord_photo_file": "Landlord Photo",
    "aadhaar_file": "Aadhaar",
    "pan_file": "PAN",
    "property_tax_file": "Property Tax",
    "passbook_file": "Passbook",
}

ALLOWED_FILE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".pdf"}

SECTION_REMARK_LABELS = {
    "site_location": "Site Location",
    "landlord_contact": "Landlord Contact Details",
    "size_rental_display": "Size, Rental & Display",
    "documents": "Photos & Documents",
    "gps_agreement": "GPS & Agreement",
}


def safe_content_disposition(filename: str, disposition: str = "inline") -> str:
    original = filename or "file"
    fallback = re.sub(r'[^A-Za-z0-9._-]+', "_", original).strip("._") or "file"
    encoded = quote(original, safe="")
    return f"{disposition}; filename=\"{fallback}\"; filename*=UTF-8''{encoded}"


def now_utc() -> datetime:
    return datetime.utcnow()


def env_bool(name: str, default: bool = False) -> bool:
    value = os.getenv(name)
    if value is None:
        return default
    return value.strip().lower() in {"1", "true", "yes", "on"}


def ensure_schema_compatibility():
    """Add new columns without deleting or rewriting existing data."""
    inspector = inspect(engine)
    if "ooh_sites" not in inspector.get_table_names():
        return
    existing = {column["name"] for column in inspector.get_columns("ooh_sites")}
    dialect = engine.dialect.name

    if dialect == "postgresql":
        definitions = {
            "size_boxes": "JSON DEFAULT '[]' NOT NULL",
            "rent_amount": "DOUBLE PRECISION DEFAULT 0 NOT NULL",
            "agreement_created": "BOOLEAN DEFAULT false NOT NULL",
            "agreement_created_at": "TIMESTAMP NULL",
            "interested_status": "VARCHAR(30) DEFAULT 'Not Selected' NOT NULL",
        }
    else:
        definitions = {
            "size_boxes": "JSON DEFAULT '[]' NOT NULL",
            "rent_amount": "FLOAT DEFAULT 0 NOT NULL",
            "agreement_created": "BOOLEAN DEFAULT 0 NOT NULL",
            "agreement_created_at": "DATETIME NULL",
            "interested_status": "VARCHAR(30) DEFAULT 'Not Selected' NOT NULL",
        }

    with engine.begin() as conn:
        for column_name, ddl in definitions.items():
            if column_name not in existing:
                conn.execute(sql_text(f"ALTER TABLE ooh_sites ADD COLUMN {column_name} {ddl}"))

        # Keep the database default aligned with the form default without overwriting valid existing choices.
        if "interested_status" in existing or "interested_status" in definitions:
            if dialect == "postgresql":
                conn.execute(sql_text("ALTER TABLE ooh_sites ALTER COLUMN interested_status SET DEFAULT 'Not Selected'"))
            conn.execute(sql_text("UPDATE ooh_sites SET interested_status = 'Not Selected' WHERE interested_status IS NULL OR interested_status NOT IN ('Not Selected', 'Interested', 'Not Interested')"))


def normalize_size_boxes(size_boxes_json: str | list | None, width_ft: float, height_ft: float, side_type: str) -> list[dict]:
    raw = None
    if isinstance(size_boxes_json, list):
        raw = size_boxes_json
    elif size_boxes_json:
        try:
            raw = json.loads(size_boxes_json)
        except json.JSONDecodeError:
            raw = None

    if not isinstance(raw, list) or not raw:
        raw = [{"label": "Single Side" if side_type == "Single" else "Side 1", "width_ft": width_ft, "height_ft": height_ft}]

    normalized = []
    for index, box in enumerate(raw, start=1):
        if not isinstance(box, dict):
            continue
        width = float(box.get("width_ft") or 0)
        height = float(box.get("height_ft") or 0)
        if width <= 0 or height <= 0:
            raise HTTPException(status_code=400, detail="Every size box must have width and height greater than zero")
        label = str(box.get("label") or f"Size {index}").strip()
        normalized.append({
            "label": label,
            "width_ft": width,
            "height_ft": height,
            "area_sqft": round(width * height, 2),
        })
    if not normalized:
        raise HTTPException(status_code=400, detail="At least one valid size box is required")
    return normalized


def get_site_for_user_or_404(site_id: int, db: Session, current_user: User) -> OOHSite:
    site = db.query(OOHSite).filter(OOHSite.id == int(site_id)).first()
    if not site:
        raise HTTPException(status_code=404, detail="Site not found")
    if current_user.role != "admin" and int(site.created_by_user_id) != int(current_user.id):
        raise HTTPException(status_code=403, detail="Not allowed")
    return site


def init_db():
    init_database()
    ensure_schema_compatibility()
    db = next(get_db())
    try:
        admin_email = os.getenv("SEED_ADMIN_EMAIL", "admin@adinn.com").strip().lower()
        admin_password = os.getenv("SEED_ADMIN_PASSWORD", "Admin@123")
        admin_name = os.getenv("SEED_ADMIN_NAME", "ADINN Admin").strip() or "ADINN Admin"

        admin = db.query(User).filter(User.email == admin_email).first()
        if not admin:
            db.add(
                User(
                    name=admin_name,
                    email=admin_email,
                    phone="",
                    password_hash=hash_password(admin_password),
                    role="admin",
                    is_active=True,
                    can_crud=True,
                    created_at=now_utc(),
                    updated_at=now_utc(),
                )
            )
            db.commit()

        if env_bool("CREATE_DEMO_EMPLOYEE", True):
            employee_email = os.getenv("DEMO_EMPLOYEE_EMAIL", "employee@adinn.com").strip().lower()
            employee_password = os.getenv("DEMO_EMPLOYEE_PASSWORD", "Employee@123")
            employee_name = os.getenv("DEMO_EMPLOYEE_NAME", "Field Employee").strip() or "Field Employee"
            employee = db.query(User).filter(User.email == employee_email).first()
            if not employee:
                db.add(
                    User(
                        name=employee_name,
                        email=employee_email,
                        phone="",
                        password_hash=hash_password(employee_password),
                        role="employee",
                        is_active=True,
                        can_crud=True,
                        created_at=now_utc(),
                        updated_at=now_utc(),
                    )
                )
                db.commit()

        # Requirement update: employees get CRUD access by default.
        db.query(User).filter(User.role == "employee", User.can_crud == False).update(  # noqa: E712
            {"can_crud": True, "updated_at": now_utc()}, synchronize_session=False
        )

        # Backfill new size-box structure for old records without deleting data.
        for legacy_site in db.query(OOHSite).all():
            if not legacy_site.size_boxes:
                legacy_site.size_boxes = normalize_size_boxes(None, legacy_site.width_ft, legacy_site.height_ft, legacy_site.side_type)
            if legacy_site.rent_amount is None:
                legacy_site.rent_amount = 0
        db.commit()
    finally:
        db.close()


@app.on_event("startup")
def startup_event():
    init_db()


def user_to_dict(user: User) -> dict:
    return {
        "id": int(user.id),
        "name": user.name or "",
        "email": user.email or "",
        "phone": user.phone,
        "role": user.role or "employee",
        "is_active": bool(user.is_active),
        "can_crud": bool(user.can_crud),
        "created_at": user.created_at or now_utc(),
    }


def normalize_role(role: str) -> str:
    clean_role = (role or "employee").strip().lower()
    if clean_role not in {"admin", "employee"}:
        raise HTTPException(status_code=400, detail="Role must be admin or employee")
    return clean_role


def count_active_admins(db: Session) -> int:
    return db.query(User).filter(User.role == "admin", User.is_active == True).count()  # noqa: E712


def build_user_updates(payload: UserUpdate, db: Session, user_id: int) -> dict:
    role = normalize_role(payload.role)
    email = str(payload.email).strip().lower()
    existing = db.query(User).filter(User.email == email, User.id != int(user_id)).first()
    if existing:
        raise HTTPException(status_code=409, detail="Email already exists")

    updates = {
        "name": payload.name.strip(),
        "email": email,
        "phone": payload.phone,
        "role": role,
        "is_active": bool(payload.is_active),
        "can_crud": True if role == "admin" else bool(payload.can_crud),
        "updated_at": now_utc(),
    }
    if payload.password and payload.password.strip():
        if len(payload.password.strip()) < 6:
            raise HTTPException(status_code=400, detail="Password must be at least 6 characters")
        updates["password_hash"] = hash_password(payload.password.strip())
    return updates


def validate_file(file: UploadFile | None):
    if file is None:
        return
    ext = Path(file.filename or "").suffix.lower()
    if ext not in ALLOWED_FILE_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Unsupported file type for {file.filename}. Allowed: jpg, jpeg, png, pdf",
        )


def save_upload_to_neon(site_id: int, upload: UploadFile | None, field_name: str, db: Session) -> str | None:
    """Store uploaded file bytes directly in Neon PostgreSQL.

    The OOH table keeps a lightweight filename marker, while the actual file
    bytes live in the stored_files table. Existing local file paths are not
    deleted; they are used as a fallback only when no DB file exists.
    """
    if upload is None:
        return None
    validate_file(upload)
    ext = Path(upload.filename or "").suffix.lower()
    safe_name = f"{field_name}_{uuid4().hex}{ext}"
    content = upload.file.read()
    if not content:
        raise HTTPException(status_code=400, detail=f"{upload.filename or field_name} is empty")

    existing = db.query(StoredFile).filter(
        StoredFile.site_id == int(site_id),
        StoredFile.field_name == field_name,
    ).first()
    if existing:
        existing.original_filename = upload.filename or safe_name
        existing.content_type = upload.content_type or "application/octet-stream"
        existing.file_size = len(content)
        existing.data = content
        existing.updated_at = now_utc()
    else:
        db.add(StoredFile(
            site_id=int(site_id),
            field_name=field_name,
            original_filename=upload.filename or safe_name,
            content_type=upload.content_type or "application/octet-stream",
            file_size=len(content),
            data=content,
            created_at=now_utc(),
            updated_at=now_utc(),
        ))
    return safe_name


def file_url(site_id: int | None, field_name: str, marker: str | None) -> str | None:
    if not marker or not site_id:
        return None
    return f"/api/sites/{site_id}/files/{field_name}"


def validate_site_payload(
    height_ft: float,
    width_ft: float,
    side_type: str,
    towards_1: str | None,
    towards_2: str | None,
    agreement_start_date: date,
    agreement_end_date: date,
):
    if height_ft <= 0 or width_ft <= 0:
        raise HTTPException(status_code=400, detail="Height and width must be greater than zero")
    if side_type == "Single" and not towards_1:
        raise HTTPException(status_code=400, detail="Towards field is required for single side")
    if side_type == "Double" and (not towards_1 or not towards_2):
        raise HTTPException(status_code=400, detail="Both towards fields are required for double side")
    if agreement_end_date < agreement_start_date:
        raise HTTPException(status_code=400, detail="Agreement end date cannot be before start date")


def parse_remarks_json(remarks_json: str | dict | None) -> dict[str, str]:
    if not remarks_json:
        return {}
    if isinstance(remarks_json, dict):
        data = remarks_json
    else:
        try:
            data = json.loads(remarks_json)
        except json.JSONDecodeError:
            return {}
    if not isinstance(data, dict):
        return {}
    return {str(key): str(value).strip() for key, value in data.items() if str(value).strip()}


def format_remarks_for_excel(remarks: dict | str | None) -> str:
    parsed = parse_remarks_json(remarks)
    return " | ".join(f"{SECTION_REMARK_LABELS.get(key, key)}: {value}" for key, value in parsed.items())


def normalize_interested_status(value: str | None) -> str:
    cleaned = (value or "Not Selected").strip()
    if cleaned not in {"Not Selected", "Interested", "Not Interested"}:
        raise HTTPException(status_code=400, detail="Interest status must be Not Selected, Interested, or Not Interested")
    return cleaned


def make_document_data(site: OOHSite | None) -> dict:
    return {
        "site_photo_uploaded": bool(site and site.site_photo_file),
        "landlord_photo_uploaded": bool(site and site.landlord_photo_file),
        "aadhaar_uploaded": bool(site and site.aadhaar_file),
        "pan_uploaded": bool(site and site.pan_file),
        "property_tax_uploaded": bool(site and site.property_tax_file),
        "passbook_uploaded": bool(site and site.passbook_file),
        "site_photo_file_url": file_url(site.id if site else None, "site_photo_file", site.site_photo_file if site else None),
        "landlord_photo_file_url": file_url(site.id if site else None, "landlord_photo_file", site.landlord_photo_file if site else None),
        "aadhaar_file_url": file_url(site.id if site else None, "aadhaar_file", site.aadhaar_file if site else None),
        "pan_file_url": file_url(site.id if site else None, "pan_file", site.pan_file if site else None),
        "property_tax_file_url": file_url(site.id if site else None, "property_tax_file", site.property_tax_file if site else None),
        "passbook_file_url": file_url(site.id if site else None, "passbook_file", site.passbook_file if site else None),
    }


def site_to_out(site: OOHSite, db: Session) -> SiteOut:
    creator = db.query(User).filter(User.id == site.created_by_user_id).first()
    return SiteOut(
        id=int(site.id),
        created_by_user_id=int(site.created_by_user_id),
        created_by_name=creator.name if creator else None,
        hoarding_location=site.hoarding_location,
        landlord_location=site.landlord_location,
        landlord_name=site.landlord_name,
        landlord_phone=site.landlord_phone,
        landlord_email=site.landlord_email,
        secondary_contact_name=site.secondary_contact_name,
        secondary_contact_phone=site.secondary_contact_phone,
        height_ft=float(site.height_ft),
        width_ft=float(site.width_ft),
        area_sqft=float(site.area_sqft),
        size_boxes=site.size_boxes or normalize_size_boxes(None, site.width_ft, site.height_ft, site.side_type),
        rental_type=site.rental_type,
        rent_amount=float(site.rent_amount or 0),
        advance_amount=float(site.advance_amount or 0),
        light_type=site.light_type,
        side_type=site.side_type,
        towards_1=site.towards_1,
        towards_2=site.towards_2,
        interested_status=site.interested_status or "Not Selected",
        latitude=site.latitude,
        longitude=site.longitude,
        agreement_tenure=site.agreement_tenure,
        agreement_start_date=site.agreement_start_date,
        agreement_end_date=site.agreement_end_date,
        agreement_created=bool(site.agreement_created),
        agreement_created_at=site.agreement_created_at,
        agreement_status="Agreement Created" if site.agreement_created else "Agreement Not Created",
        remarks=parse_remarks_json(site.remarks),
        documents=make_document_data(site),
        created_at=site.created_at or now_utc(),
        updated_at=site.updated_at or now_utc(),
    )


def can_manage_site(current_user: User, site: OOHSite) -> bool:
    if current_user.role == "admin":
        return True
    if site.agreement_created:
        return False
    return bool(current_user.can_crud) and int(site.created_by_user_id) == int(current_user.id)


def require_site_crud_permission(current_user: User, site: OOHSite):
    if not can_manage_site(current_user, site):
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="CRUD is blocked for this employee, this record belongs to another user, or agreement is already created.",
        )


def apply_document_uploads(
    site: OOHSite,
    db: Session,
    site_photo_file: UploadFile | None = None,
    landlord_photo_file: UploadFile | None = None,
    aadhaar_file: UploadFile | None = None,
    pan_file: UploadFile | None = None,
    property_tax_file: UploadFile | None = None,
    passbook_file: UploadFile | None = None,
):
    uploaded_files = {
        "site_photo_file": save_upload_to_neon(site.id, site_photo_file, "site_photo_file", db),
        "landlord_photo_file": save_upload_to_neon(site.id, landlord_photo_file, "landlord_photo_file", db),
        "aadhaar_file": save_upload_to_neon(site.id, aadhaar_file, "aadhaar_file", db),
        "pan_file": save_upload_to_neon(site.id, pan_file, "pan_file", db),
        "property_tax_file": save_upload_to_neon(site.id, property_tax_file, "property_tax_file", db),
        "passbook_file": save_upload_to_neon(site.id, passbook_file, "passbook_file", db),
    }
    for field_name, stored_marker in uploaded_files.items():
        if stored_marker:
            setattr(site, field_name, stored_marker)
    site.updated_at = now_utc()


def apply_site_fields(
    site: OOHSite,
    hoarding_location: str,
    landlord_location: str,
    landlord_name: str,
    landlord_phone: str,
    landlord_email: str | None,
    secondary_contact_name: str | None,
    secondary_contact_phone: str | None,
    height_ft: float,
    width_ft: float,
    size_boxes_json: str | None,
    rental_type: str,
    rent_amount: float,
    advance_amount: float,
    light_type: str,
    side_type: str,
    towards_1: str | None,
    towards_2: str | None,
    interested_status: str | None,
    latitude: float | None,
    longitude: float | None,
    agreement_tenure: str,
    agreement_start_date: date,
    agreement_end_date: date,
    remarks_json: str | None,
):
    site.hoarding_location = hoarding_location.strip()
    site.landlord_location = landlord_location.strip()
    site.landlord_name = landlord_name.strip()
    site.landlord_phone = landlord_phone.strip()
    site.landlord_email = landlord_email.strip() if landlord_email else None
    site.secondary_contact_name = secondary_contact_name.strip() if secondary_contact_name else None
    site.secondary_contact_phone = secondary_contact_phone.strip() if secondary_contact_phone else None
    normalized_boxes = normalize_size_boxes(size_boxes_json, width_ft, height_ft, side_type)
    site.size_boxes = normalized_boxes
    first_box = normalized_boxes[0]
    site.width_ft = float(first_box["width_ft"])
    site.height_ft = float(first_box["height_ft"])
    site.area_sqft = round(sum(float(box["area_sqft"]) for box in normalized_boxes), 2)
    site.rental_type = rental_type
    site.rent_amount = float(rent_amount or 0)
    site.advance_amount = float(advance_amount or 0)
    site.light_type = light_type
    site.side_type = side_type
    site.towards_1 = towards_1
    site.towards_2 = towards_2 if side_type == "Double" else None
    site.interested_status = normalize_interested_status(interested_status)
    site.latitude = latitude
    site.longitude = longitude
    site.agreement_tenure = agreement_tenure
    site.agreement_start_date = agreement_start_date
    site.agreement_end_date = agreement_end_date
    site.remarks = parse_remarks_json(remarks_json)
    site.updated_at = now_utc()


@app.get("/api/health")
def health():
    ping_database()
    return {"status": "ok", "database": "neon-postgresql"}


@app.post("/api/auth/login", response_model=TokenResponse)
def login(payload: LoginRequest, db: Session = Depends(get_db)):
    user = db.query(User).filter(User.email == str(payload.email).lower()).first()
    if not user or not verify_password(payload.password, user.password_hash):
        raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="Invalid email or password")
    if not user.is_active:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="User is inactive")
    return TokenResponse(access_token=create_token(user), user=user_to_dict(user))


@app.get("/api/users/me", response_model=UserOut)
def me(current_user: User = Depends(get_current_user)):
    return user_to_dict(current_user)


@app.post("/api/users", response_model=UserOut)
def create_user(
    payload: UserCreate,
    db: Session = Depends(get_db),
    _: User = Depends(require_admin),
):
    role = normalize_role(payload.role)
    existing = db.query(User).filter(User.email == str(payload.email).strip().lower()).first()
    if existing:
        raise HTTPException(status_code=409, detail="Email already exists")
    if len(payload.password.strip()) < 6:
        raise HTTPException(status_code=400, detail="Password must be at least 6 characters")
    user = User(
        name=payload.name.strip(),
        email=str(payload.email).strip().lower(),
        phone=payload.phone,
        password_hash=hash_password(payload.password.strip()),
        role=role,
        is_active=True,
        can_crud=bool(payload.can_crud) if role == "employee" else True,
        created_at=now_utc(),
        updated_at=now_utc(),
    )
    db.add(user)
    try:
        db.commit()
    except IntegrityError:
        db.rollback()
        raise HTTPException(status_code=409, detail="Email already exists")
    db.refresh(user)
    return user_to_dict(user)


@app.get("/api/users", response_model=list[UserOut])
def list_users(db: Session = Depends(get_db), _: User = Depends(require_admin)):
    users = db.query(User).order_by(User.created_at.desc()).all()
    return [user_to_dict(user) for user in users]


@app.patch("/api/users/{user_id}/crud-permission", response_model=UserOut)
def update_user_crud_permission(
    user_id: int,
    payload: CrudPermissionUpdate,
    db: Session = Depends(get_db),
    _: User = Depends(require_admin),
):
    user = db.query(User).filter(User.id == int(user_id)).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    user.can_crud = True if user.role == "admin" else bool(payload.can_crud)
    user.updated_at = now_utc()
    db.commit()
    db.refresh(user)
    return user_to_dict(user)


@app.put("/api/users/{user_id}", response_model=UserOut)
def update_user(
    user_id: int,
    payload: UserUpdate,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    user = db.query(User).filter(User.id == int(user_id)).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    updates = build_user_updates(payload, db, int(user_id))

    if int(user_id) == int(current_user.id):
        if not updates.get("is_active", True):
            raise HTTPException(status_code=400, detail="You cannot deactivate your own admin account")
        if updates.get("role") != "admin":
            raise HTTPException(status_code=400, detail="You cannot remove admin role from your own account")

    if user.role == "admin" and (updates.get("role") != "admin" or not updates.get("is_active", True)) and count_active_admins(db) <= 1:
        raise HTTPException(status_code=400, detail="At least one active admin must remain")

    for key, value in updates.items():
        setattr(user, key, value)
    try:
        db.commit()
    except IntegrityError:
        db.rollback()
        raise HTTPException(status_code=409, detail="Email already exists")
    db.refresh(user)
    return user_to_dict(user)


@app.delete("/api/users/{user_id}")
def delete_user(
    user_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_admin),
):
    user = db.query(User).filter(User.id == int(user_id)).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    if int(user_id) == int(current_user.id):
        raise HTTPException(status_code=400, detail="You cannot delete your own admin account")
    if user.role == "admin" and count_active_admins(db) <= 1:
        raise HTTPException(status_code=400, detail="At least one active admin must remain")
    if db.query(OOHSite).filter(OOHSite.created_by_user_id == int(user_id)).count() > 0:
        user.is_active = False
        user.can_crud = False
        user.updated_at = now_utc()
        db.commit()
        return {"message": "User has existing OOH records, so the account was deactivated instead of permanently deleted."}
    db.delete(user)
    db.commit()
    return {"message": "User deleted successfully"}


@app.post("/api/sites", response_model=SiteOut)
def create_site(
    hoarding_location: str = Form(...),
    landlord_location: str = Form(...),
    landlord_name: str = Form(...),
    landlord_phone: str = Form(...),
    landlord_email: Optional[str] = Form(None),
    secondary_contact_name: Optional[str] = Form(None),
    secondary_contact_phone: Optional[str] = Form(None),
    height_ft: float = Form(...),
    width_ft: float = Form(...),
    size_boxes_json: Optional[str] = Form(None),
    rental_type: str = Form(...),
    rent_amount: float = Form(0),
    advance_amount: float = Form(0),
    light_type: str = Form(...),
    side_type: str = Form(...),
    towards_1: Optional[str] = Form(None),
    towards_2: Optional[str] = Form(None),
    interested_status: str = Form("Not Selected"),
    latitude: Optional[float] = Form(None),
    longitude: Optional[float] = Form(None),
    agreement_tenure: str = Form(...),
    agreement_start_date: date = Form(...),
    agreement_end_date: date = Form(...),
    remarks_json: Optional[str] = Form(None),
    create_agreement: bool = Form(False),
    site_photo_file: UploadFile | None = File(None),
    landlord_photo_file: UploadFile | None = File(None),
    aadhaar_file: UploadFile | None = File(None),
    pan_file: UploadFile | None = File(None),
    property_tax_file: UploadFile | None = File(None),
    passbook_file: UploadFile | None = File(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    validate_site_payload(height_ft, width_ft, side_type, towards_1, towards_2, agreement_start_date, agreement_end_date)
    site = OOHSite(created_by_user_id=current_user.id)
    apply_site_fields(
        site,
        hoarding_location,
        landlord_location,
        landlord_name,
        landlord_phone,
        landlord_email,
        secondary_contact_name,
        secondary_contact_phone,
        height_ft,
        width_ft,
        size_boxes_json,
        rental_type,
        rent_amount,
        advance_amount,
        light_type,
        side_type,
        towards_1,
        towards_2,
        interested_status,
        latitude,
        longitude,
        agreement_tenure,
        agreement_start_date,
        agreement_end_date,
        remarks_json,
    )
    # Agreement creation is intentionally allowed only from the Records section
    # through POST /api/sites/{site_id}/agreement, never during initial site creation.
    site.agreement_created = False
    site.agreement_created_at = None
    site.created_at = now_utc()
    db.add(site)
    db.flush()
    apply_document_uploads(site, db, site_photo_file, landlord_photo_file, aadhaar_file, pan_file, property_tax_file, passbook_file)
    db.commit()
    db.refresh(site)
    return site_to_out(site, db)


@app.put("/api/sites/{site_id}", response_model=SiteOut)
def update_site(
    site_id: int,
    hoarding_location: str = Form(...),
    landlord_location: str = Form(...),
    landlord_name: str = Form(...),
    landlord_phone: str = Form(...),
    landlord_email: Optional[str] = Form(None),
    secondary_contact_name: Optional[str] = Form(None),
    secondary_contact_phone: Optional[str] = Form(None),
    height_ft: float = Form(...),
    width_ft: float = Form(...),
    size_boxes_json: Optional[str] = Form(None),
    rental_type: str = Form(...),
    rent_amount: float = Form(0),
    advance_amount: float = Form(0),
    light_type: str = Form(...),
    side_type: str = Form(...),
    towards_1: Optional[str] = Form(None),
    towards_2: Optional[str] = Form(None),
    interested_status: str = Form("Not Selected"),
    latitude: Optional[float] = Form(None),
    longitude: Optional[float] = Form(None),
    agreement_tenure: str = Form(...),
    agreement_start_date: date = Form(...),
    agreement_end_date: date = Form(...),
    remarks_json: Optional[str] = Form(None),
    site_photo_file: UploadFile | None = File(None),
    landlord_photo_file: UploadFile | None = File(None),
    aadhaar_file: UploadFile | None = File(None),
    pan_file: UploadFile | None = File(None),
    property_tax_file: UploadFile | None = File(None),
    passbook_file: UploadFile | None = File(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    site = db.query(OOHSite).filter(OOHSite.id == int(site_id)).first()
    if not site:
        raise HTTPException(status_code=404, detail="Site not found")
    require_site_crud_permission(current_user, site)
    validate_site_payload(height_ft, width_ft, side_type, towards_1, towards_2, agreement_start_date, agreement_end_date)
    apply_site_fields(
        site,
        hoarding_location,
        landlord_location,
        landlord_name,
        landlord_phone,
        landlord_email,
        secondary_contact_name,
        secondary_contact_phone,
        height_ft,
        width_ft,
        size_boxes_json,
        rental_type,
        rent_amount,
        advance_amount,
        light_type,
        side_type,
        towards_1,
        towards_2,
        interested_status,
        latitude,
        longitude,
        agreement_tenure,
        agreement_start_date,
        agreement_end_date,
        remarks_json,
    )
    apply_document_uploads(site, db, site_photo_file, landlord_photo_file, aadhaar_file, pan_file, property_tax_file, passbook_file)
    db.commit()
    db.refresh(site)
    return site_to_out(site, db)


@app.delete("/api/sites/{site_id}")
def delete_site(site_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    site = db.query(OOHSite).filter(OOHSite.id == int(site_id)).first()
    if not site:
        raise HTTPException(status_code=404, detail="Site not found")
    require_site_crud_permission(current_user, site)
    db.query(StoredFile).filter(StoredFile.site_id == int(site_id)).delete(synchronize_session=False)
    db.delete(site)
    db.commit()
    shutil.rmtree(UPLOAD_DIR / f"site_{site_id}", ignore_errors=True)
    return {"message": "OOH site deleted successfully"}


@app.get("/api/sites", response_model=list[SiteOut])
def list_sites(db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    query = db.query(OOHSite)
    if current_user.role != "admin":
        query = query.filter(OOHSite.created_by_user_id == int(current_user.id))
    sites = query.order_by(OOHSite.created_at.desc()).all()
    return [site_to_out(site, db) for site in sites]

@app.get("/api/sites/{site_id}/files/{field_name}")
def get_site_file(
    site_id: int,
    field_name: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    if field_name not in FILE_FIELD_LABELS:
        raise HTTPException(status_code=404, detail="File field not found")
    site = get_site_for_user_or_404(site_id, db, current_user)
    marker = getattr(site, field_name, None)
    if not marker:
        raise HTTPException(status_code=404, detail="File not uploaded")

    stored = db.query(StoredFile).filter(
        StoredFile.site_id == int(site_id),
        StoredFile.field_name == field_name,
    ).first()
    if stored and stored.data:
        filename = stored.original_filename or f"{field_name}.bin"
        return Response(
            content=stored.data,
            media_type=stored.content_type or "application/octet-stream",
            headers={"Content-Disposition": safe_content_disposition(filename)},
        )

    # Fallback for old local uploads created before DB-backed file storage.
    legacy_path = UPLOAD_DIR / str(marker)
    if legacy_path.exists() and legacy_path.is_file():
        return FileResponse(legacy_path)
    raise HTTPException(status_code=404, detail="File is referenced, but the stored file was not found")


@app.post("/api/sites/{site_id}/agreement", response_model=SiteOut)
def create_site_agreement(site_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    site = get_site_for_user_or_404(site_id, db, current_user)
    if site.agreement_created:
        raise HTTPException(status_code=400, detail="Agreement already created for this OOH site")
    if current_user.role != "admin" and int(site.created_by_user_id) != int(current_user.id):
        raise HTTPException(status_code=403, detail="Only the owner or admin can create agreement")
    site.agreement_created = True
    site.agreement_created_at = now_utc()
    site.updated_at = now_utc()
    db.commit()
    db.refresh(site)
    return site_to_out(site, db)


def site_export_rows(site: OOHSite, db: Session) -> list[tuple[str, str]]:
    creator = db.query(User).filter(User.id == site.created_by_user_id).first()
    size_summary = "; ".join(
        f"{box.get('label') or 'Size'}: {box.get('width_ft')} ft W × {box.get('height_ft')} ft H = {box.get('area_sqft')} sqft"
        for box in (site.size_boxes or normalize_size_boxes(None, site.width_ft, site.height_ft, site.side_type))
    )
    return [
        ("OOH ID", str(site.id)),
        ("Agreement Status", "Agreement Created" if site.agreement_created else "Agreement Not Created"),
        ("Uploaded On", (site.created_at or now_utc()).strftime("%Y-%m-%d %H:%M")),
        ("Last Updated", (site.updated_at or now_utc()).strftime("%Y-%m-%d %H:%M")),
        ("Entered By", creator.name if creator else ""),
        ("Hoarding Location", site.hoarding_location or ""),
        ("Landlord Location", site.landlord_location or ""),
        ("Landlord Name", site.landlord_name or ""),
        ("Landlord Phone", site.landlord_phone or ""),
        ("Landlord Email", site.landlord_email or ""),
        ("Secondary Contact", f"{site.secondary_contact_name or ''} {site.secondary_contact_phone or ''}".strip()),
        ("Size Details", size_summary),
        ("Total Area", f"{site.area_sqft} sqft"),
        ("Rental Type", site.rental_type or ""),
        ("Rent", f"Rs {site.rent_amount or 0:,.0f}"),
        ("Advance", f"Rs {site.advance_amount or 0:,.0f}"),
        ("Light", site.light_type or ""),
        ("Side", site.side_type or ""),
        ("Interest Status", site.interested_status or "Not Selected"),
        ("Towards 1", site.towards_1 or ""),
        ("Towards 2", site.towards_2 or ""),
        ("GPS", f"{site.latitude or ''}, {site.longitude or ''}".strip(', ')),
        ("Agreement Tenure", site.agreement_tenure or ""),
        ("Agreement Dates", f"{site.agreement_start_date} to {site.agreement_end_date}"),
        ("Remarks", format_remarks_for_excel(site.remarks)),
    ]


@app.get("/api/sites/{site_id}/export/docx")
def export_site_docx(site_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    site = get_site_for_user_or_404(site_id, db, current_user)
    document = Document()
    document.add_heading(f"ADINN OOH Site Details #{site.id}", level=1)
    document.add_paragraph("Generated from ADINN OOH Site Tracker")
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Field"
    table.rows[0].cells[1].text = "Details"
    for label, value in site_export_rows(site, db):
        row = table.add_row().cells
        row[0].text = label
        row[1].text = value
    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=ooh_site_{site.id}.docx"},
    )


@app.get("/api/sites/{site_id}/export/pdf")
def export_site_pdf(site_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    site = get_site_for_user_or_404(site_id, db, current_user)
    output = io.BytesIO()
    doc = SimpleDocTemplate(output, pagesize=A4, rightMargin=28, leftMargin=28, topMargin=30, bottomMargin=30)
    styles = getSampleStyleSheet()
    elements = [Paragraph(f"ADINN OOH Site Details #{site.id}", styles["Title"]), Spacer(1, 10)]
    data = [[Paragraph("<b>Field</b>", styles["Normal"]), Paragraph("<b>Details</b>", styles["Normal"])] ]
    for label, value in site_export_rows(site, db):
        data.append([Paragraph(str(label), styles["Normal"]), Paragraph(str(value).replace("\n", "<br/>") or "-", styles["Normal"])])
    table = Table(data, colWidths=[135, 365])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#101828")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#D0D5DD")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("BACKGROUND", (0, 1), (0, -1), colors.HexColor("#FFF1EE")),
    ]))
    elements.append(table)
    doc.build(elements)
    output.seek(0)
    return StreamingResponse(
        output,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=ooh_site_{site.id}.pdf"},
    )


def build_sites_excel_response(sites: list[OOHSite], db: Session, filename: str = "adinn_ooh_sites.xlsx"):
    wb = Workbook()
    ws = wb.active
    ws.title = "OOH Sites"
    headers = [
        "ID", "Entered By", "Hoarding Location", "Landlord Location", "Landlord Name",
        "Landlord Phone", "Landlord Email", "Secondary Name", "Secondary Phone",
        "Width ft", "Height ft", "Area sqft", "Size Boxes", "Rental Type", "Rent Rs", "Advance Rs",
        "Light", "Side", "Interest Status", "Towards 1", "Towards 2", "Latitude", "Longitude",
        "Agreement Tenure", "Start Date", "End Date", "Agreement Status", "Agreement Created At", "Site Photo", "Landlord Photo", "Aadhaar", "PAN", "Property Tax",
        "Passbook", "Remarks", "Created At"
    ]
    ws.append(headers)
    for site in sites:
        creator = db.query(User).filter(User.id == site.created_by_user_id).first()
        ws.append([
            site.id,
            creator.name if creator else "",
            site.hoarding_location,
            site.landlord_location,
            site.landlord_name,
            site.landlord_phone,
            site.landlord_email or "",
            site.secondary_contact_name or "",
            site.secondary_contact_phone or "",
            site.width_ft,
            site.height_ft,
            site.area_sqft,
            "; ".join(f"{box.get('label')}: {box.get('width_ft')}W x {box.get('height_ft')}H = {box.get('area_sqft')} sqft" for box in (site.size_boxes or [])),
            site.rental_type,
            site.rent_amount or 0,
            site.advance_amount,
            site.light_type,
            site.side_type,
            site.interested_status or "Not Selected",
            site.towards_1 or "",
            site.towards_2 or "",
            site.latitude or "",
            site.longitude or "",
            site.agreement_tenure,
            site.agreement_start_date.isoformat() if site.agreement_start_date else "",
            site.agreement_end_date.isoformat() if site.agreement_end_date else "",
            "Agreement Created" if site.agreement_created else "Agreement Not Created",
            site.agreement_created_at.strftime("%Y-%m-%d %H:%M") if site.agreement_created_at else "",
            "Uploaded" if site.site_photo_file else "Not Uploaded",
            "Uploaded" if site.landlord_photo_file else "Not Uploaded",
            "Uploaded" if site.aadhaar_file else "Not Uploaded",
            "Uploaded" if site.pan_file else "Not Uploaded",
            "Uploaded" if site.property_tax_file else "Not Uploaded",
            "Uploaded" if site.passbook_file else "Not Uploaded",
            format_remarks_for_excel(site.remarks),
            (site.created_at or now_utc()).strftime("%Y-%m-%d %H:%M"),
        ])
    for column_cells in ws.columns:
        max_length = max(len(str(cell.value)) if cell.value is not None else 0 for cell in column_cells)
        ws.column_dimensions[column_cells[0].column_letter].width = min(max(max_length + 2, 12), 45)
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": safe_content_disposition(filename, "attachment")},
    )


@app.get("/api/sites/export/excel")
def export_sites_excel(db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    query = db.query(OOHSite)
    if current_user.role != "admin":
        query = query.filter(OOHSite.created_by_user_id == int(current_user.id))
    sites = query.order_by(OOHSite.created_at.desc()).all()
    return build_sites_excel_response(sites, db, "adinn_ooh_sites.xlsx")


@app.post("/api/sites/export/excel/selected")
def export_selected_sites_excel(payload: dict, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    raw_ids = payload.get("site_ids") or []
    if not isinstance(raw_ids, list) or not raw_ids:
        raise HTTPException(status_code=400, detail="Select at least one hoarding to export")
    try:
        ids = [int(value) for value in raw_ids]
    except (TypeError, ValueError):
        raise HTTPException(status_code=400, detail="Invalid selected hoarding IDs")
    unique_ids = list(dict.fromkeys(ids))
    sites = db.query(OOHSite).filter(OOHSite.id.in_(unique_ids)).order_by(OOHSite.created_at.desc()).all()
    if not sites:
        raise HTTPException(status_code=404, detail="No selected hoardings found")
    return build_sites_excel_response(sites, db, "adinn_ooh_selected_sites.xlsx")

@app.get("/api/sites/{site_id}", response_model=SiteOut)
def get_site(site_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    site = get_site_for_user_or_404(site_id, db, current_user)
    return site_to_out(site, db)
